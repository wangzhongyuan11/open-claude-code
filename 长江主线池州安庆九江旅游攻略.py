#!/usr/bin/env python3
"""
长江主线池州安庆九江旅游攻略 - Word文档生成脚本
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def set_chinese_font(paragraph):
    """设置中文段落字体"""
    for run in paragraph.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)


def add_heading(doc, text, level=1):
    """添加标题"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 1:
            run.font.size = Pt(24)
        elif level == 2:
            run.font.size = Pt(18)
        elif level == 3:
            run.font.size = Pt(14)
    return heading


def add_paragraph(doc, text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """添加段落"""
    para = doc.add_paragraph()
    para.alignment = alignment
    run = para.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    return para


def add_list(doc, items, level=0):
    """添加列表"""
    for item in items:
        para = doc.add_paragraph(style=f'List Paragraph')
        run = para.add_run(item)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)


def main():
    # 创建文档
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # 封面
    add_heading(doc, '长江主线池州-安庆-九江旅游攻略', level=1)
    add_paragraph(doc, '', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, '路线：池州 → 安庆 → 九江', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, f'更新时间：{datetime.now().strftime("%Y年%m月%d日")}', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, '', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, '长江是中国第一大河，流经安徽、江西等省份，沿岸城市拥有丰富的历史文化和自然风光。', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, '本攻略以长江为主线，串联池州、安庆、九江三座城市，带您领略长江流域的独特魅力。', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    doc.add_page_break()
    
    # 目录
    add_heading(doc, '目录', level=1)
    doc.add_paragraph('一、池州旅游攻略', style='List Number')
    doc.add_paragraph('    1.1 城市概况', style='List Number 2')
    doc.add_paragraph('    1.2 核心景点与游览路线', style='List Number 2')
    doc.add_paragraph('    1.3 特色美食', style='List Number 2')
    doc.add_paragraph('二、安庆旅游攻略', style='List Number')
    doc.add_paragraph('    2.1 城市概况', style='List Number 2')
    doc.add_paragraph('    2.2 核心景点与游览路线', style='List Number 2')
    doc.add_paragraph('    2.3 特色美食', style='List Number 2')
    doc.add_paragraph('三、九江旅游攻略', style='List Number')
    doc.add_paragraph('    3.1 城市概况', style='List Number 2')
    doc.add_paragraph('    3.2 核心景点与游览路线', style='List Number 2')
    doc.add_paragraph('    3.3 特色美食', style='List Number 2')
    doc.add_paragraph('四、交通与住宿建议', style='List Number')
    doc.add_paragraph('五、注意事项', style='List Number')
    
    doc.add_page_break()
    
    # 池州旅游攻略
    add_heading(doc, '一、池州旅游攻略', level=1)
    
    add_heading(doc, '1.1 城市概况', level=2)
    add_paragraph(doc, '池州，别名"千载诗人地"，位于安徽省西南部，长江下游南岸。这里背倚九华山，怀抱秋浦河，是长江流域的重要节点城市。')
    add_paragraph(doc, '历史文化：池州历史悠久，自唐武德四年（621年）设州，距今已有1400多年历史。作为"中国佛教四大名山"之一的九华山所在地，池州拥有深厚的佛教文化底蕴。')
    add_paragraph(doc, '与长江的关系：池州地处长江南岸，长江流经池州境内约162公里，是长江经济带的重要组成部分。长江不仅为池州提供了便利的水运交通，更塑造了池州独特的滨江风光。')
    
    add_heading(doc, '1.2 核心景点与游览路线', level=2)
    add_paragraph(doc, '推荐行程（2天1晚）：', bold=True)
    add_paragraph(doc, '第一天：九华山风景区 → 肉身宝殿 → 化城寺 → 天台峰')
    add_paragraph(doc, '第二天：杏花村 → 秋浦河漂流 → 池州长江大桥')
    
    add_paragraph(doc, '核心景点介绍：', bold=True)
    add_list(doc, [
        '九华山风景区：中国佛教四大名山之一，以地藏菩萨道场闻名，有"东南第一山"之称。',
        '肉身宝殿：九华山最著名的建筑之一，供奉着地藏菩萨的肉身。',
        '化城寺：九华山的开山祖寺，是九华山佛教文化的发源地。',
        '天台峰：九华山的主峰之一，海拔1306米，是观赏日出日落的最佳地点。',
        '杏花村：唐代诗人杜牧《清明》诗中"牧童遥指杏花村"的所在地，是一处集田园风光与文化古迹于一体的景区。',
        '秋浦河漂流：秋浦河是池州的母亲河，漂流河段全长约15公里，两岸风光旖旎。',
        '池州长江大桥：连接池州市区与枞阳县的重要通道，桥面可俯瞰长江壮丽景色。'
    ])
    
    add_heading(doc, '1.3 特色美食', level=2)
    add_paragraph(doc, '池州美食以徽菜为主，注重原汁原味，口感醇厚。推荐美食：', bold=True)
    add_list(doc, [
        '池州臭鳜鱼：徽州名菜之一，肉质鲜嫩，闻起来臭，吃起来香。',
        '九华素宴：以素食为主，食材多为山珍野菌，营养丰富。',
        '石台一品锅：将各种食材分层放入锅中蒸熟，味道鲜美。',
        '东至米饺：以大米为原料，馅料丰富，口感软糯。',
        '青阳年糕：用糯米制作而成，口感细腻，可蒸可煮。'
    ])
    
    doc.add_page_break()
    
    # 安庆旅游攻略
    add_heading(doc, '二、安庆旅游攻略', level=1)
    
    add_heading(doc, '2.1 城市概况', level=2)
    add_paragraph(doc, '安庆，别名"宜城"，位于安徽省西南部，长江下游北岸。这里是皖西南的中心城市，也是长江流域的重要文化名城。')
    add_paragraph(doc, '历史文化：安庆历史悠久，自南宋嘉定十年（1217年）建城，距今已有800多年历史。安庆是桐城派故里和黄梅戏之乡，曾长期作为安徽省会，享有"文化之邦"、"戏剧之乡"的美誉。')
    add_paragraph(doc, '与长江的关系：安庆地处长江北岸，长江流经安庆境内约243公里，是长江经济带的重要节点城市。长江不仅为安庆提供了便利的水运交通，更塑造了安庆独特的滨江风光。')
    
    add_heading(doc, '2.2 核心景点与游览路线', level=2)
    add_paragraph(doc, '推荐行程（2天1晚）：', bold=True)
    add_paragraph(doc, '第一天：天柱山风景区 → 三祖寺 → 炼丹湖')
    add_paragraph(doc, '第二天：迎江寺 → 振风塔 → 安庆长江大桥 → 安庆老城')
    
    add_paragraph(doc, '核心景点介绍：', bold=True)
    add_list(doc, [
        '天柱山风景区：国家5A级旅游景区，以雄奇灵秀的自然风光著称，有"江淮第一山"之称。',
        '三祖寺：天柱山的开山祖寺，是佛教禅宗的发源地之一。',
        '炼丹湖：中国第三大高山人工湖，湖面清澈见底，周围群山环绕。',
        '迎江寺：安庆市最著名的佛教寺庙之一，始建于北宋开宝七年（974年）。',
        '振风塔：安庆市的标志性建筑之一，始建于明隆庆四年（1570年），高约60米。',
        '安庆长江大桥：连接安庆市与池州市的重要通道，桥面可俯瞰长江壮丽景色。',
        '安庆老城：保存完好的明清古城，是安庆历史文化的缩影。'
    ])
    
    add_heading(doc, '2.3 特色美食', level=2)
    add_paragraph(doc, '安庆美食以徽菜为主，口味偏重，注重火候。推荐美食：', bold=True)
    add_list(doc, [
        '安庆炒面：以面条为主料，配以各种蔬菜和肉类，口感丰富。',
        '安庆包子：以面粉为原料，馅料丰富，口感松软。',
        '安庆馄饨：以面粉为原料，馅料细腻，汤头鲜美。',
        '怀宁贡糕：用糯米制作而成，口感细腻，甜而不腻。',
        '潜山舒席：以竹为原料，编织而成，可作为餐具使用。'
    ])
    
    doc.add_page_break()
    
    # 九江旅游攻略
    add_heading(doc, '三、九江旅游攻略', level=1)
    
    add_heading(doc, '3.1 城市概况', level=2)
    add_paragraph(doc, '九江，别名"浔阳"，位于江西省北部，长江中游南岸。这里是江西的北大门，也是长江流域的重要港口城市。')
    add_paragraph(doc, '历史文化：九江历史悠久，自秦代设九江郡，距今已有2200多年历史。九江是中国历史文化名城，拥有丰富的历史文化遗产和自然风光。')
    add_paragraph(doc, '与长江的关系：九江地处长江中游南岸，长江流经九江境内约151公里，是长江经济带的重要节点城市。九江港是长江流域的重要港口之一，水运交通十分便利。')
    
    add_heading(doc, '3.2 核心景点与游览路线', level=2)
    add_paragraph(doc, '推荐行程（2天1晚）：', bold=True)
    add_paragraph(doc, '第一天：庐山风景区 → 庐山瀑布 → 锦绣谷 → 仙人洞')
    add_paragraph(doc, '第二天：浔阳楼 → 琵琶亭 → 锁江楼 → 九江长江大桥')
    
    add_paragraph(doc, '核心景点介绍：', bold=True)
    add_list(doc, [
        '庐山风景区：国家5A级旅游景区，以雄奇险秀的自然风光和深厚的文化底蕴著称，有"匡庐奇秀甲天下"之称。',
        '庐山瀑布：李白诗中"飞流直下三千尺，疑是银河落九天"的所在地，是庐山最著名的景点之一。',
        '锦绣谷：庐山的一处峡谷景观，谷中怪石嶙峋，风景秀丽。',
        '仙人洞：庐山的一处天然石洞，传说是吕洞宾修炼成仙的地方。',
        '浔阳楼：唐代诗人白居易《琵琶行》诗中"浔阳江头夜送客"的所在地，是一处集文化古迹与滨江风光于一体的景区。',
        '琵琶亭：纪念白居易《琵琶行》的建筑，亭内有琵琶女的雕像。',
        '锁江楼：始建于明代万历年间，是九江的标志性建筑之一，楼前有一座铁牛，相传是为了镇住长江洪水。',
        '九江长江大桥：连接九江市与湖北省黄梅县的重要通道，是长江上最长的公路铁路两用桥之一。'
    ])
    
    add_heading(doc, '3.3 特色美食', level=2)
    add_paragraph(doc, '九江美食以赣菜为主，口味偏重，注重调味。推荐美食：', bold=True)
    add_list(doc, [
        '九江炒粉：以米粉为主料，配以各种蔬菜和肉类，口感丰富。',
        '九江酱鸭：以鸭肉为主料，用多种香料腌制而成，味道鲜美。',
        '九江茶饼：以面粉为原料，馅料丰富，口感酥脆。',
        '湖口酒糟鱼：以鲤鱼为主料，用酒糟腌制而成，味道鲜美。',
        '修水哨子：以米粉为原料，馅料丰富，口感松软。'
    ])
    
    doc.add_page_break()
    
    # 交通与住宿建议
    add_heading(doc, '四、交通与住宿建议', level=1)
    
    add_paragraph(doc, '交通建议：', bold=True)
    add_list(doc, [
        '飞机：池州九华山机场、安庆天柱山机场、九江庐山机场均有航班通往国内主要城市。',
        '高铁：池州站、安庆站、九江站均有高铁通往国内主要城市。',
        '汽车：长江流域的城市之间有高速公路相连，交通便利。',
        '水运：长江是中国重要的水运通道，可乘坐游轮游览长江沿岸城市。'
    ])
    
    add_paragraph(doc, '住宿建议：', bold=True)
    add_list(doc, [
        '池州市区：推荐住宿在九华山脚下或池州市区，交通便利。',
        '安庆市区：推荐住宿在迎江区或大观区，靠近景点。',
        '九江市区：推荐住宿在浔阳区或庐山区，靠近景点。'
    ])
    
    # 注意事项
    add_heading(doc, '五、注意事项', level=1)
    add_list(doc, [
        '长江流域气候湿润，夏季炎热，冬季寒冷，建议根据季节携带合适的衣物。',
        '长江流域的景点多为山区或滨江地区，建议穿着舒适的鞋子。',
        '长江流域的饮食以辣为主，建议携带肠胃药。',
        '长江流域的水运交通便利，但需注意安全。',
        '长江流域的景点多为自然景观，建议携带防晒霜、帽子等防晒用品。'
    ])
    
    # 保存文档
    output_dir = 'output/doc'
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    output_path = os.path.join(output_dir, '长江主线池州-安庆-九江旅游攻略.docx')
    doc.save(output_path)
    print(f'旅游攻略已生成：{output_path}')


if __name__ == '__main__':
    main()
